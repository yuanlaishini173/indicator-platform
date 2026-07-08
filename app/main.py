from __future__ import annotations

import base64
import hashlib
import html
import io
import json
import os
import re
import secrets
import shutil
import socket
import sqlite3
import subprocess
import sys
import time
import urllib.error
import urllib.request
from contextlib import contextmanager
from datetime import datetime
from pathlib import Path
from typing import Any

from fastapi import FastAPI, File, Form, HTTPException, Query, Request, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, PlainTextResponse, RedirectResponse, Response
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel


ROOT = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT / "data"
UPLOAD_DIR = ROOT / "uploads"
STATIC_DIR = Path(__file__).resolve().parent / "static"
DB_PATH = DATA_DIR / "platform.db"
REPORT_DIR = DATA_DIR / "reports"
PUBLIC_SHARE_URL_FILE = ROOT / "public-share-url.txt"
PUBLIC_SHARE_STATUS_FILE = DATA_DIR / "public-share-status.json"
GITHUB_PAGES_SHARE_URL = os.getenv(
    "GITHUB_PAGES_SHARE_URL",
    "",
).strip()
CLOUDFLARED_EXE = ROOT / "tools" / "cloudflared.exe"
CLOUDFLARED_LOG = DATA_DIR / "cloudflared-public.log"
CLOUDFLARED_OUT_LOG = DATA_DIR / "cloudflared-public.out.log"
CLOUDFLARED_ERR_LOG = DATA_DIR / "cloudflared-public.err.log"

RUNTIME_SITE_PACKAGES = (
    Path.home()
    / ".cache"
    / "codex-runtimes"
    / "codex-primary-runtime"
    / "dependencies"
    / "python"
    / "Lib"
    / "site-packages"
)
if RUNTIME_SITE_PACKAGES.exists():
    sys.path.append(str(RUNTIME_SITE_PACKAGES))

try:
    import pandas as pd
except Exception:
    pd = None


class ManualMetricPayload(BaseModel):
    month: str | None = None
    actual_value: float | None = None


class UploadUpdatePayload(BaseModel):
    filename: str | None = None


class OpenReportFolderPayload(BaseModel):
    filename: str | None = None


class RulePayload(BaseModel):
    indicator_pattern: str
    category: str
    direction: str = "higher_better"
    warning_threshold: float = 80
    done_threshold: float = 100


APP_TITLE = "指标完成情况平台"
DEFAULT_MONTH = datetime.now().strftime("%Y-%m")
HISTORICAL_PERIOD_RE = re.compile(r"^HIST-(20\d{2})-(20\d{2})$")
METRIC_VALUE_KEYS = [
    "total_certificates",
    "total_people",
    "advanced_worker_plus",
    "technician_plus",
    "special_technician",
    "chief_technician",
]

STATUS_LABELS = {
    "done": "已完成",
    "warning": "接近完成",
    "behind": "未完成",
    "unknown": "待判断",
}

TOTAL_DISTRICTS = {"合计", "总计", "全市合计", "南京市"}
IGNORED_SOURCE_DISTRICTS = {"南京化学工业园", "江北新区（旭东南路）", "南京市人才库房(二)", "高层人才", "南京市"}
DISTRICT_MERGE_MAP = {
    "南京经济技术开发区": "栖霞区",
    "南京江北新区": "直管区",
    "江北新区": "直管区",
    "直管区": "直管区",
    "浦口区": "浦口区",
}
NORTH_BANK_COMPONENTS = ("直管区", "浦口区")

FORMAL_INDICATORS = [
    {
        "name": "新增取得职业资格证书或职业技能等级证书人数（万人次）",
        "target": 5.8,
        "unit": "万人次",
        "source": "total_certificates",
        "note": "取证数合计 / 10000。",
    },
    {
        "name": "新增取得高级工以上职业资格证书或职业技能等级证书人数（人次）",
        "target": 20000,
        "unit": "人次",
        "source": "advanced_worker_plus",
        "note": "三级+二级+一级+特级技师+首席技师取证数。",
    },
    {
        "name": "新增取得技师以上职业资格证书或职业技能等级证书人数（人次）",
        "target": 1700,
        "unit": "人次",
        "source": "technician_plus",
        "note": "二级+一级+特级技师+首席技师取证数。",
    },
    {
        "name": "新增专业技能人才取得职业技能等级人数（人）",
        "target": 580,
        "unit": "人",
        "source": "professional_skill",
        "note": "手工填入或专项表汇总。",
    },
    {
        "name": "新增取得人工智能类职业资格证书或职业技能等级证书人数（人次）",
        "target": 24000,
        "unit": "人次",
        "source": "ai_certificates",
        "note": "人工智能类专项表合计行取证数。",
    },
]
FORMAL_INDICATOR_ORDER = {item["name"]: index for index, item in enumerate(FORMAL_INDICATORS, start=1)}

DISTRICT_TARGETS = {
    FORMAL_INDICATORS[0]["name"]: {
        "合计": 5.8,
        "北岸": 0.9,
        "直管区": 0.54,
        "浦口区": 0.36,
        "玄武区": 0.3,
        "秦淮区": 0.4,
        "建邺区": 0.38,
        "鼓楼区": 0.6,
        "栖霞区": 0.85,
        "雨花台区": 0.33,
        "江宁区": 0.9,
        "六合区": 0.28,
        "溧水区": 0.28,
        "高淳区": 0.28,
    },
    FORMAL_INDICATORS[1]["name"]: {
        "合计": 20000,
        "北岸": 3700,
        "直管区": 2000,
        "浦口区": 1700,
        "玄武区": 1000,
        "秦淮区": 1500,
        "建邺区": 1300,
        "鼓楼区": 3300,
        "栖霞区": 3900,
        "雨花台区": 1400,
        "江宁区": 2700,
        "六合区": 400,
        "溧水区": 400,
        "高淳区": 400,
    },
    FORMAL_INDICATORS[2]["name"]: {"合计": 1700},
    FORMAL_INDICATORS[3]["name"]: {"合计": 580},
    FORMAL_INDICATORS[4]["name"]: {"合计": 24000},
}

SKILL_COLUMNS = {
    "district": ["地区", "区域", "区县", "区市", "单位"],
    "total": ["取证数合计", "取证总数", "合计"],
    "level3": ["三级取证数", "三级"],
    "level2": ["二级取证数", "二级"],
    "level1": ["一级取证数", "一级"],
    "special": ["特级技师取证数", "特技技师取证数", "特级技师", "特技技师"],
    "chief": ["首席技师取证数", "首席技师"],
}


@contextmanager
def db() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    try:
        yield conn
        conn.commit()
    finally:
        conn.close()


def now_iso() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def rows_to_dicts(rows: list[sqlite3.Row]) -> list[dict[str, Any]]:
    return [dict(row) for row in rows]


def init_db() -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    with db() as conn:
        conn.executescript(
            """
            CREATE TABLE IF NOT EXISTS uploads (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                filename TEXT NOT NULL,
                stored_path TEXT,
                month TEXT NOT NULL,
                uploaded_at TEXT NOT NULL,
                sheet_names TEXT NOT NULL DEFAULT '[]',
                row_count INTEGER NOT NULL DEFAULT 0,
                status TEXT NOT NULL DEFAULT 'parsed'
            );

            CREATE TABLE IF NOT EXISTS metrics (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                upload_id INTEGER,
                month TEXT NOT NULL,
                district TEXT NOT NULL,
                indicator_code TEXT NOT NULL,
                indicator_name TEXT NOT NULL,
                category TEXT NOT NULL,
                target_value REAL,
                actual_value REAL,
                completion_rate REAL,
                unit TEXT,
                direction TEXT NOT NULL DEFAULT 'higher_better',
                status TEXT NOT NULL DEFAULT 'unknown',
                source_sheet TEXT,
                source_row INTEGER,
                notes TEXT,
                created_at TEXT NOT NULL,
                FOREIGN KEY(upload_id) REFERENCES uploads(id)
            );

            CREATE TABLE IF NOT EXISTS ai_district_metrics (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                upload_id INTEGER,
                month TEXT NOT NULL,
                district TEXT NOT NULL,
                actual_value REAL NOT NULL,
                unit TEXT NOT NULL DEFAULT '人次',
                source_sheet TEXT,
                created_at TEXT NOT NULL,
                FOREIGN KEY(upload_id) REFERENCES uploads(id)
            );

            CREATE TABLE IF NOT EXISTS certificate_level_metrics (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                upload_id INTEGER,
                month TEXT NOT NULL,
                district TEXT NOT NULL,
                special_technician REAL NOT NULL DEFAULT 0,
                chief_technician REAL NOT NULL DEFAULT 0,
                unit TEXT NOT NULL DEFAULT '人次',
                source_sheet TEXT,
                created_at TEXT NOT NULL,
                FOREIGN KEY(upload_id) REFERENCES uploads(id)
            );

            CREATE TABLE IF NOT EXISTS indicator_rules (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                indicator_pattern TEXT NOT NULL,
                category TEXT NOT NULL,
                direction TEXT NOT NULL DEFAULT 'higher_better',
                warning_threshold REAL NOT NULL DEFAULT 80,
                done_threshold REAL NOT NULL DEFAULT 100,
                created_at TEXT NOT NULL
            );

            CREATE INDEX IF NOT EXISTS idx_metrics_month ON metrics(month);
            CREATE INDEX IF NOT EXISTS idx_metrics_district ON metrics(district);
            CREATE INDEX IF NOT EXISTS idx_metrics_indicator ON metrics(indicator_name);
            CREATE INDEX IF NOT EXISTS idx_metrics_status ON metrics(status);
            CREATE INDEX IF NOT EXISTS idx_ai_district_month ON ai_district_metrics(month);
            CREATE INDEX IF NOT EXISTS idx_level_metrics_month ON certificate_level_metrics(month);
            """
        )


def normalize_month(value: str | None) -> str:
    text = str(value or "").strip()
    historical = historical_period_code_from_text(text)
    if historical:
        return historical
    match = re.search(r"(20\d{2})[-年./]?(\d{1,2})", text)
    if match:
        return f"{match.group(1)}-{int(match.group(2)):02d}"
    return DEFAULT_MONTH


def normalize_year(value: str) -> int:
    year = int(value)
    return year + 2000 if year < 100 else year


def historical_period_code_from_text(value: str | None) -> str | None:
    text = str(value or "").strip()
    direct = HISTORICAL_PERIOD_RE.match(text)
    if direct:
        start, end = int(direct.group(1)), int(direct.group(2))
        return f"HIST-{start}-{end}" if start <= end else None
    match = re.search(r"(\d{2,4})年?\s*(?:至|到|-|—|--|~)\s*(\d{2,4})年", text)
    if not match:
        return None
    start = normalize_year(match.group(1))
    end = normalize_year(match.group(2))
    if start >= 2000 and end >= start:
        return f"HIST-{start}-{end}"
    return None


def is_historical_period(value: str | None) -> bool:
    return bool(HISTORICAL_PERIOD_RE.match(str(value or "")))


def historical_period_label(value: str) -> str:
    match = HISTORICAL_PERIOD_RE.match(str(value or ""))
    if not match:
        return str(value or "")
    return f"{match.group(1)}-{match.group(2)}年累计"


def period_sort_value(value: str | None) -> int:
    text = str(value or "")
    historical = HISTORICAL_PERIOD_RE.match(text)
    if historical:
        return int(historical.group(2)) * 12 + 12
    match = re.match(r"^(20\d{2})-(\d{2})$", text)
    if match:
        return int(match.group(1)) * 12 + int(match.group(2))
    return 0


def infer_month_from_text(value: str | None) -> str | None:
    if not value:
        return None
    text = str(value)
    historical = historical_period_code_from_text(text)
    if historical:
        return historical
    match = re.search(r"(\d{2,4})年?(\d{1,2})个?月", text)
    if match:
        year = int(match.group(1))
        if year < 100:
            year += 2000
        return f"{year}-{int(match.group(2)):02d}"
    match = re.search(r"(20\d{2})[-_./](\d{1,2})", text)
    if match:
        return f"{match.group(1)}-{int(match.group(2)):02d}"
    return None


def latest_month() -> str:
    with db() as conn:
        rows = [row["month"] for row in conn.execute("SELECT DISTINCT month FROM metrics").fetchall()]
    return sorted(rows, key=period_sort_value, reverse=True)[0] if rows else DEFAULT_MONTH


def period_label(month: str) -> str:
    if is_historical_period(month):
        return historical_period_label(month)
    year, month_num = month.split("-")
    number = int(month_num)
    return f"{year}年1月累计" if number <= 1 else f"{year}年1-{number}月累计"


def month_label(month: str) -> str:
    if is_historical_period(month):
        return historical_period_label(month)
    year, month_num = month.split("-")
    return f"{year}年{int(month_num)}月"


def month_to_int(month: str) -> int:
    if is_historical_period(month):
        return period_sort_value(month)
    year, month_num = normalize_month(month).split("-")
    return int(year) * 12 + int(month_num)


def previous_month(month: str) -> str:
    year, month_num = normalize_month(month).split("-")
    year_num = int(year)
    number = int(month_num)
    if number == 1:
        return f"{year_num - 1}-12"
    return f"{year_num}-{number - 1:02d}"


def normalize_time_scope(value: str | None) -> str:
    if value == "since2021":
        return "since2021"
    return "range" if value == "range" else "cumulative"


def resolve_time_period(month: str | None, start_month: str | None = None, time_scope: str | None = None) -> dict[str, Any]:
    end_month = normalize_month(month) if month else latest_month()
    scope = normalize_time_scope(time_scope)
    start = normalize_month(start_month) if start_month else end_month
    if is_historical_period(end_month) or is_historical_period(start):
        scope = "cumulative"
    if month_to_int(start) > month_to_int(end_month):
        start = end_month
    if scope == "since2021":
        start = "HIST-2021-2025"
        baseline = ""
        label = f"2021年至{month_label(end_month)}累计"
    elif scope != "range":
        start = ""
        baseline = ""
        label = period_label(end_month)
    else:
        baseline = previous_month(start)
        label = month_label(end_month) if start == end_month else f"{month_label(start)}-{month_label(end_month)}"
    return {"month": end_month, "start_month": start, "baseline_month": baseline, "time_scope": scope, "label": label}


def can_use_zero_baseline(period: dict[str, Any]) -> bool:
    return period.get("time_scope") == "range" and str(period.get("start_month") or "").endswith("-01")


def normalize_report_period(value: str | None) -> str:
    return value if value in {"month", "quarter"} else "month"


def normalize_report_scope(value: str | None) -> str:
    return value if value in {"single", "cumulative"} else "cumulative"


def make_indicator_code(name: str) -> str:
    return hashlib.md5(name.encode("utf-8")).hexdigest()[:8]


def clean_number(value: Any) -> float | None:
    if value is None:
        return None
    if isinstance(value, (int, float)) and not pd.isna(value):
        return float(value)
    text = str(value).strip()
    if not text or text.lower() == "nan":
        return None
    text = text.replace(",", "").replace("，", "").replace("%", "")
    match = re.search(r"-?\d+(?:\.\d+)?", text)
    return float(match.group(0)) if match else None


def normalize_header(value: Any) -> str:
    text = str(value or "").strip()
    text = re.sub(r"\s+", "", text)
    text = text.replace("\n", "").replace("\r", "")
    return text


def normalize_output_district(value: Any) -> str | None:
    name = str(value or "").strip()
    if not name or name.lower() == "nan":
        return None
    name = re.sub(r"\s+", "", name)
    name = name.replace("其中：", "").replace("其中:", "").replace("……", "").replace("...", "")
    if name in IGNORED_SOURCE_DISTRICTS:
        return None
    if name in TOTAL_DISTRICTS:
        return "合计"
    return DISTRICT_MERGE_MAP.get(name, name)


def target_for_indicator(indicator_name: str, district: str) -> float | None:
    value = DISTRICT_TARGETS.get(indicator_name, {}).get(district)
    return float(value) if value is not None else None


def calculate_completion(target: float | None, actual: float | None) -> float | None:
    if target in (None, 0) or actual is None:
        return None
    return round(float(actual) / float(target) * 100, 2)


def completion_status(rate: float | None) -> str:
    if rate is None:
        return "unknown"
    if rate >= 100:
        return "done"
    if rate >= 80:
        return "warning"
    return "behind"


def format_value(value: Any, unit: str | None = "") -> str:
    number = clean_number(value)
    if number is None:
        return "-"
    text = str(int(number)) if float(number).is_integer() else f"{number:.2f}".rstrip("0").rstrip(".")
    return f"{text}{unit or ''}"


def city_metric_row(indicator: dict[str, Any], actual: float | None, month: str, upload_id: int | None, sheet: str, notes: str) -> dict[str, Any]:
    name = indicator["name"]
    target = None if is_historical_period(month) else target_for_indicator(name, "合计")
    rate = calculate_completion(target, actual)
    row_notes = f"{notes} 历史累计期仅统计取证数量，不计算2026年度目标完成率。" if is_historical_period(month) else notes
    return {
        "upload_id": upload_id,
        "month": month,
        "district": "合计",
        "indicator_code": make_indicator_code(name),
        "indicator_name": name,
        "category": "职业技能证书",
        "target_value": target,
        "actual_value": actual,
        "completion_rate": rate,
        "unit": indicator["unit"],
        "direction": "higher_better",
        "status": completion_status(rate),
        "source_sheet": sheet,
        "source_row": None,
        "notes": row_notes,
        "created_at": now_iso(),
    }


def metric_from_values(district: str, indicator: dict[str, Any], actual: float | None, month: str, upload_id: int | None, sheet: str, notes: str) -> dict[str, Any]:
    name = indicator["name"]
    target = None if is_historical_period(month) else target_for_indicator(name, district)
    rate = calculate_completion(target, actual)
    row_notes = f"{notes} 历史累计期仅统计取证数量，不计算2026年度目标完成率。" if is_historical_period(month) else notes
    return {
        "upload_id": upload_id,
        "month": month,
        "district": district,
        "indicator_code": make_indicator_code(name),
        "indicator_name": name,
        "category": "职业技能证书",
        "target_value": target,
        "actual_value": actual,
        "completion_rate": rate,
        "unit": indicator["unit"],
        "direction": "higher_better",
        "status": completion_status(rate),
        "source_sheet": sheet,
        "source_row": None,
        "notes": row_notes,
        "created_at": now_iso(),
    }


def find_columns(df: "pd.DataFrame") -> dict[str, str]:
    normalized = {normalize_header(column): column for column in df.columns}
    found: dict[str, str] = {}
    for key, aliases in SKILL_COLUMNS.items():
        for alias in aliases:
            if alias in normalized:
                found[key] = normalized[alias]
                break
        if key not in found:
            for normalized_name, original in normalized.items():
                if any(alias in normalized_name for alias in aliases):
                    found[key] = original
                    break
    return found


def values_from_row(row: Any, cols: dict[str, str]) -> dict[str, float]:
    def get(key: str) -> float:
        return clean_number(row.get(cols.get(key, ""))) or 0.0

    total = get("total")
    special = get("special")
    chief = get("chief")
    advanced = get("level3") + get("level2") + get("level1") + special + chief
    technician = get("level2") + get("level1") + special + chief
    return {
        "total_certificates": round(total / 10000, 4),
        "total_people": round(total, 2),
        "advanced_worker_plus": round(advanced, 2),
        "technician_plus": round(technician, 2),
        "special_technician": round(special, 2),
        "chief_technician": round(chief, 2),
    }


def aggregate_skill_table(path: Path) -> tuple[dict[str, dict[str, float]], list[str]]:
    if pd is None:
        raise HTTPException(status_code=500, detail="当前运行环境缺少 pandas/openpyxl，暂不能解析表格")
    excel = pd.read_excel(path, sheet_name=None, dtype=object)
    aggregated: dict[str, dict[str, float]] = {}
    sheet_names: list[str] = []
    for sheet_name, raw_df in excel.items():
        df = raw_df.dropna(how="all")
        if df.empty:
            continue
        cols = find_columns(df)
        if "district" not in cols or "total" not in cols:
            continue
        sheet_names.append(sheet_name)
        for _, row in df.iterrows():
            district = normalize_output_district(row.get(cols["district"]))
            if not district:
                continue
            values = values_from_row(row, cols)
            bucket = aggregated.setdefault(district, {key: 0.0 for key in values})
            for key, value in values.items():
                bucket[key] += value
    if "合计" not in aggregated:
        city = {key: 0.0 for key in METRIC_VALUE_KEYS}
        for district, values in aggregated.items():
            if district not in {"北岸"}:
                for key in city:
                    city[key] += values.get(key, 0.0)
        if any(city.values()):
            aggregated["合计"] = city
    if all(component in aggregated for component in NORTH_BANK_COMPONENTS):
        north = {key: aggregated["直管区"].get(key, 0.0) + aggregated["浦口区"].get(key, 0.0) for key in aggregated["直管区"]}
        aggregated["北岸"] = north
    return aggregated, sheet_names


def is_ai_file(filename: str) -> bool:
    return "人工智能" in filename


def is_professional_file(filename: str) -> bool:
    return "专业人才" in filename or "专业技能人才" in filename


def level_rows_from_values(values: dict[str, dict[str, float]], month: str, upload_id: int | None, sheet: str) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for district, district_values in values.items():
        rows.append(
            {
                "upload_id": upload_id,
                "month": month,
                "district": district,
                "special_technician": round(district_values.get("special_technician", 0.0), 2),
                "chief_technician": round(district_values.get("chief_technician", 0.0), 2),
                "unit": "人次",
                "source_sheet": sheet,
                "created_at": now_iso(),
            }
        )
    return rows


def parse_spreadsheet(path: Path, month: str, upload_id: int, source_name: str) -> tuple[list[dict[str, Any]], list[str], list[dict[str, Any]], list[dict[str, Any]]]:
    values, sheets = aggregate_skill_table(path)
    rows: list[dict[str, Any]] = []
    ai_rows: list[dict[str, Any]] = []
    level_rows: list[dict[str, Any]] = []
    sheet = "、".join(sheets) if sheets else "sheet"

    if is_ai_file(source_name):
        indicator = FORMAL_INDICATORS[4]
        city_actual = values.get("合计", {}).get("total_people")
        rows.append(city_metric_row(indicator, city_actual, month, upload_id, sheet, indicator["note"]))
        for district, district_values in values.items():
            if district == "合计":
                continue
            ai_rows.append(
                {
                    "upload_id": upload_id,
                    "month": month,
                    "district": district,
                    "actual_value": round(district_values.get("total_people", 0.0), 2),
                    "unit": "人次",
                    "source_sheet": sheet,
                    "created_at": now_iso(),
                }
            )
        return rows, sheets, ai_rows, level_rows

    if is_professional_file(source_name):
        indicator = FORMAL_INDICATORS[3]
        city_actual = values.get("合计", {}).get("total_people")
        rows.append(city_metric_row(indicator, city_actual, month, upload_id, sheet, indicator["note"]))
        return rows, sheets, [], level_rows

    level_rows = level_rows_from_values(values, month, upload_id, sheet)
    for district, district_values in values.items():
        if district == "合计":
            for indicator in FORMAL_INDICATORS[:3]:
                rows.append(metric_from_values(district, indicator, district_values.get(indicator["source"]), month, upload_id, sheet, indicator["note"]))
            continue
        for indicator in FORMAL_INDICATORS[:2]:
            rows.append(metric_from_values(district, indicator, district_values.get(indicator["source"]), month, upload_id, sheet, indicator["note"]))
    return rows, sheets, [], level_rows


def insert_metrics(rows: list[dict[str, Any]]) -> None:
    if not rows:
        return
    keys = sorted({(row["month"], row["district"], row["indicator_name"]) for row in rows})
    fields = [
        "upload_id",
        "month",
        "district",
        "indicator_code",
        "indicator_name",
        "category",
        "target_value",
        "actual_value",
        "completion_rate",
        "unit",
        "direction",
        "status",
        "source_sheet",
        "source_row",
        "notes",
        "created_at",
    ]
    placeholders = ",".join(["?"] * len(fields))
    with db() as conn:
        conn.executemany("DELETE FROM metrics WHERE month = ? AND district = ? AND indicator_name = ?", keys)
        conn.executemany(
            f"INSERT INTO metrics ({','.join(fields)}) VALUES ({placeholders})",
            [[row.get(field) for field in fields] for row in rows],
        )


def insert_ai_district_metrics(rows: list[dict[str, Any]]) -> None:
    if not rows:
        return
    month = rows[0]["month"]
    fields = ["upload_id", "month", "district", "actual_value", "unit", "source_sheet", "created_at"]
    placeholders = ",".join(["?"] * len(fields))
    with db() as conn:
        conn.execute("DELETE FROM ai_district_metrics WHERE month = ?", (month,))
        conn.executemany(
            f"INSERT INTO ai_district_metrics ({','.join(fields)}) VALUES ({placeholders})",
            [[row.get(field) for field in fields] for row in rows],
        )


def insert_certificate_level_metrics(rows: list[dict[str, Any]]) -> None:
    if not rows:
        return
    month = rows[0]["month"]
    fields = ["upload_id", "month", "district", "special_technician", "chief_technician", "unit", "source_sheet", "created_at"]
    placeholders = ",".join(["?"] * len(fields))
    with db() as conn:
        conn.execute("DELETE FROM certificate_level_metrics WHERE month = ?", (month,))
        conn.executemany(
            f"INSERT INTO certificate_level_metrics ({','.join(fields)}) VALUES ({placeholders})",
            [[row.get(field) for field in fields] for row in rows],
        )


def backfill_certificate_level_metrics() -> None:
    with db() as conn:
        existing_months = {row["month"] for row in conn.execute("SELECT DISTINCT month FROM certificate_level_metrics").fetchall()}
        upload_rows = rows_to_dicts(
            conn.execute(
                """
                SELECT id, filename, stored_path, month
                FROM uploads
                WHERE status = 'parsed'
                ORDER BY month, id
                """
            ).fetchall()
        )
    latest_skill_upload: dict[str, dict[str, Any]] = {}
    for row in upload_rows:
        filename = row.get("filename") or ""
        if is_ai_file(filename) or is_professional_file(filename):
            continue
        path_text = str(row.get("stored_path") or "").strip()
        if not path_text:
            continue
        path = Path(path_text)
        if not path.exists():
            continue
        latest_skill_upload[row["month"]] = row

    for month, row in latest_skill_upload.items():
        if month in existing_months:
            continue
        try:
            values, sheets = aggregate_skill_table(Path(row["stored_path"]))
            sheet = "、".join(sheets) if sheets else "sheet"
            insert_certificate_level_metrics(level_rows_from_values(values, month, row["id"], sheet))
        except Exception:
            # 回填失败不影响平台启动；后续重新上传该月表格会自动写入。
            continue


def create_upload_record(filename: str, stored_path: str, month: str) -> int:
    with db() as conn:
        cursor = conn.execute(
            """
            INSERT INTO uploads (filename, stored_path, month, uploaded_at, sheet_names, row_count, status)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (filename, stored_path, month, now_iso(), "[]", 0, "uploading"),
        )
    return int(cursor.lastrowid)


def update_upload_record(upload_id: int, row_count: int, sheet_names: list[str], status: str = "parsed") -> None:
    with db() as conn:
        conn.execute(
            "UPDATE uploads SET row_count = ?, sheet_names = ?, status = ? WHERE id = ?",
            (row_count, json.dumps(sheet_names, ensure_ascii=False), status, upload_id),
        )


def get_metric_rows(
    month: str,
    district: str | None = None,
    category: str | None = None,
    status: str | None = None,
    q: str | None = None,
    limit: int = 500,
) -> list[dict[str, Any]]:
    clauses = ["month = ?"]
    params: list[Any] = [month]
    if district:
        clauses.append("district = ?")
        params.append(district)
    if category:
        clauses.append("category = ?")
        params.append(category)
    if status:
        clauses.append("status = ?")
        params.append(status)
    if q:
        clauses.append("(indicator_name LIKE ? OR district LIKE ? OR category LIKE ?)")
        like = f"%{q}%"
        params.extend([like, like, like])
    where = " AND ".join(clauses)
    with db() as conn:
        rows = rows_to_dicts(
            conn.execute(
                f"""
                SELECT * FROM metrics
                WHERE {where}
                ORDER BY district = '合计' DESC, district, indicator_name
                LIMIT ?
                """,
                [*params, limit],
            ).fetchall()
        )
    return rows


def round_metric_actual(value: float, unit: str | None) -> float:
    return round(value, 4) if unit == "万人次" else round(value, 2)


def with_recalculated_actual(row: dict[str, Any], actual: float | None, source_note: str | None = None) -> dict[str, Any]:
    item = dict(row)
    item["actual_value"] = actual
    item["completion_rate"] = calculate_completion(item.get("target_value"), actual)
    item["status"] = completion_status(item.get("completion_rate"))
    if source_note:
        item["source_sheet"] = source_note
    return item


def since_2021_metric_rows(end_rows: list[dict[str, Any]], period: dict[str, Any]) -> list[dict[str, Any]]:
    historical_rows = get_metric_rows("HIST-2021-2025", limit=5000)
    rows_by_key: dict[tuple[str, str], dict[str, Any]] = {}
    for row in historical_rows:
        rows_by_key[(row["district"], row["indicator_name"])] = dict(row)
    for row in end_rows:
        rows_by_key.setdefault((row["district"], row["indicator_name"]), dict(row))

    end_map = {(row["district"], row["indicator_name"]): row for row in end_rows}
    historical_map = {(row["district"], row["indicator_name"]): row for row in historical_rows}
    rows: list[dict[str, Any]] = []
    for key, base_row in rows_by_key.items():
        current = clean_number(end_map.get(key, {}).get("actual_value"))
        historical = clean_number(historical_map.get(key, {}).get("actual_value"))
        if current is None and historical is None:
            actual = None
        else:
            actual = round_metric_actual((current or 0) + (historical or 0), base_row.get("unit"))
        end_row = end_map.get(key, {})
        target = end_row.get("target_value")
        if target is None:
            target = target_for_indicator(base_row["indicator_name"], base_row["district"])
        rate = calculate_completion(target, actual)
        item = dict(base_row)
        item["month"] = period["month"]
        item["target_value"] = target
        item["actual_value"] = actual
        item["completion_rate"] = rate
        item["status"] = completion_status(rate)
        item["source_sheet"] = f"{period['label']}（2021-2025历史累计 + {period_label(period['month'])}）"
        item["notes"] = f"{item.get('notes') or ''} 2021年以来口径统计历史累计与本年累计合计，完成率按2026年度目标计算。".strip()
        rows.append(item)
    return rows


def metric_rows_for_period(month: str, start_month: str | None = None, time_scope: str | None = None) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    period = resolve_time_period(month, start_month, time_scope)
    end_rows = get_metric_rows(period["month"], limit=5000)
    if period["time_scope"] == "since2021":
        return since_2021_metric_rows(end_rows, period), period
    if period["time_scope"] != "range":
        return end_rows, period

    baseline_rows = get_metric_rows(period["baseline_month"], limit=5000)
    baseline_map = {(row["district"], row["indicator_name"]): row for row in baseline_rows}
    period_rows: list[dict[str, Any]] = []
    for row in end_rows:
        current = clean_number(row.get("actual_value"))
        if current is None:
            actual = None
        else:
            baseline_row = baseline_map.get((row["district"], row["indicator_name"]))
            baseline_value = clean_number(baseline_row.get("actual_value")) if baseline_row else None
            if baseline_value is None and not can_use_zero_baseline(period):
                actual = None
            else:
                diff = current - (baseline_value or 0)
                if abs(diff) < 0.000001:
                    diff = 0
                actual = round_metric_actual(max(0, diff), row.get("unit"))
        period_rows.append(with_recalculated_actual(row, actual, f"{period['label']}（累计快照相减）"))
    return period_rows, period


def filter_metric_rows(
    rows: list[dict[str, Any]],
    district: str | None = None,
    category: str | None = None,
    status: str | None = None,
    q: str | None = None,
    limit: int = 500,
) -> list[dict[str, Any]]:
    filtered = rows
    if district:
        filtered = [row for row in filtered if row.get("district") == district]
    if category:
        filtered = [row for row in filtered if row.get("category") == category]
    if status:
        filtered = [row for row in filtered if row.get("status") == status]
    if q:
        needle = q.strip().lower()
        filtered = [
            row
            for row in filtered
            if needle in str(row.get("indicator_name", "")).lower()
            or needle in str(row.get("district", "")).lower()
            or needle in str(row.get("category", "")).lower()
        ]
    filtered.sort(key=lambda item: (item.get("district") != "合计", item.get("district") or "", FORMAL_INDICATOR_ORDER.get(item.get("indicator_name"), 99)))
    return filtered[:limit]


def ai_district_rows_for_period(month: str, start_month: str | None = None, time_scope: str | None = None) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    period = resolve_time_period(month, start_month, time_scope)
    with db() as conn:
        end_rows = rows_to_dicts(
            conn.execute(
                """
                SELECT district, actual_value, unit, source_sheet
                FROM ai_district_metrics
                WHERE month = ?
                """,
                (period["month"],),
            ).fetchall()
        )
        baseline_rows = rows_to_dicts(
            conn.execute(
                """
                SELECT district, actual_value
                FROM ai_district_metrics
                WHERE month = ?
                """,
                (period["baseline_month"],),
            ).fetchall()
        ) if period["time_scope"] == "range" else []

    baseline_map = {row["district"]: clean_number(row.get("actual_value")) or 0 for row in baseline_rows}
    rows: list[dict[str, Any]] = []
    for row in end_rows:
        item = dict(row)
        if period["time_scope"] == "since2021":
            item["source_sheet"] = f"{period['label']}（人工智能历史专项数据未提供，仅含{period_label(period['month'])}）"
        if period["time_scope"] == "range":
            current = clean_number(item.get("actual_value"))
            if current is None or (item["district"] not in baseline_map and not can_use_zero_baseline(period)):
                item["actual_value"] = None
            else:
                diff = current - baseline_map.get(item["district"], 0)
                item["actual_value"] = round(max(0, diff), 2)
            item["source_sheet"] = f"{period['label']}（累计快照相减）"
        rows.append(item)
    rows.sort(key=lambda item: (-float(item.get("actual_value") or 0), item.get("district") or ""))
    return rows, period


def certificate_level_rows_for_period(month: str, start_month: str | None = None, time_scope: str | None = None) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    period = resolve_time_period(month, start_month, time_scope)
    with db() as conn:
        end_rows = rows_to_dicts(
            conn.execute(
                """
                SELECT district, special_technician, chief_technician, unit, source_sheet
                FROM certificate_level_metrics
                WHERE month = ?
                """,
                (period["month"],),
            ).fetchall()
        )
        historical_rows = rows_to_dicts(
            conn.execute(
                """
                SELECT district, special_technician, chief_technician, unit, source_sheet
                FROM certificate_level_metrics
                WHERE month = 'HIST-2021-2025'
                """
            ).fetchall()
        ) if period["time_scope"] == "since2021" else []
        baseline_rows = rows_to_dicts(
            conn.execute(
                """
                SELECT district, special_technician, chief_technician
                FROM certificate_level_metrics
                WHERE month = ?
                """,
                (period["baseline_month"],),
            ).fetchall()
        ) if period["time_scope"] == "range" else []

    if period["time_scope"] == "since2021":
        end_map = {row["district"]: row for row in end_rows}
        historical_map = {row["district"]: row for row in historical_rows}
        districts = sorted(set(end_map) | set(historical_map))
        rows: list[dict[str, Any]] = []
        for district in districts:
            end = end_map.get(district, {})
            historical = historical_map.get(district, {})
            special = (clean_number(end.get("special_technician")) or 0) + (clean_number(historical.get("special_technician")) or 0)
            chief = (clean_number(end.get("chief_technician")) or 0) + (clean_number(historical.get("chief_technician")) or 0)
            item = {
                "district": district,
                "special_technician": round(special, 2),
                "chief_technician": round(chief, 2),
                "unit": end.get("unit") or historical.get("unit") or "人次",
                "source_sheet": f"{period['label']}（2021-2025历史累计 + {period_label(period['month'])}）",
                "total_technician": round(special + chief, 2),
            }
            rows.append(item)
        rows.sort(key=lambda item: (item.get("district") != "合计", -float(item.get("total_technician") or 0), item.get("district") or ""))
        return rows, period

    baseline_map = {row["district"]: row for row in baseline_rows}
    rows: list[dict[str, Any]] = []
    for row in end_rows:
        item = dict(row)
        if period["time_scope"] == "range":
            baseline = baseline_map.get(item["district"], {})
            if not baseline and not can_use_zero_baseline(period):
                item["special_technician"] = None
                item["chief_technician"] = None
            else:
                special = (clean_number(item.get("special_technician")) or 0) - (clean_number(baseline.get("special_technician")) or 0)
                chief = (clean_number(item.get("chief_technician")) or 0) - (clean_number(baseline.get("chief_technician")) or 0)
                item["special_technician"] = round(max(0, special), 2)
                item["chief_technician"] = round(max(0, chief), 2)
            item["source_sheet"] = f"{period['label']}（累计快照相减）"
        special_value = clean_number(item.get("special_technician"))
        chief_value = clean_number(item.get("chief_technician"))
        item["total_technician"] = None if special_value is None and chief_value is None else round((special_value or 0) + (chief_value or 0), 2)
        rows.append(item)
    rows.sort(key=lambda item: (item.get("district") != "合计", -float(item.get("total_technician") or 0), item.get("district") or ""))
    return rows, period


def empty_city_metric(indicator: dict[str, Any], month: str) -> dict[str, Any]:
    target = None if is_historical_period(month) else target_for_indicator(indicator["name"], "合计")
    return {
        "id": None,
        "upload_id": None,
        "month": month,
        "district": "合计",
        "indicator_code": make_indicator_code(indicator["name"]),
        "indicator_name": indicator["name"],
        "category": "职业技能证书",
        "target_value": target,
        "actual_value": None,
        "completion_rate": None,
        "unit": indicator["unit"],
        "direction": "higher_better",
        "status": "unknown",
        "source_sheet": "",
        "source_row": None,
        "notes": indicator["note"],
        "created_at": "",
    }


def aggregate_from_rows(rows: list[dict[str, Any]], period: dict[str, Any]) -> dict[str, Any]:
    for indicator in FORMAL_INDICATORS:
        if not any(row["district"] == "合计" and row["indicator_name"] == indicator["name"] for row in rows):
            rows.append(empty_city_metric(indicator, period["month"]))

    city = [row for row in rows if row["district"] == "合计"]
    city.sort(key=lambda item: FORMAL_INDICATOR_ORDER.get(item["indicator_name"], 99))
    detail_rows = [row for row in rows if row["district"] != "合计"]

    district_map: dict[str, list[dict[str, Any]]] = {}
    for row in detail_rows:
        if row["indicator_name"] in {FORMAL_INDICATORS[0]["name"], FORMAL_INDICATORS[1]["name"]}:
            district_map.setdefault(row["district"], []).append(row)

    districts: list[dict[str, Any]] = []
    for district, items in district_map.items():
        tasks = sorted(items, key=lambda item: FORMAL_INDICATOR_ORDER.get(item["indicator_name"], 99))
        rates = [item["completion_rate"] for item in tasks if item.get("completion_rate") is not None]
        done = sum(1 for item in tasks if item.get("status") == "done")
        warning = sum(1 for item in tasks if item.get("status") == "warning")
        behind = sum(1 for item in tasks if item.get("status") == "behind")
        unknown = sum(1 for item in tasks if item.get("status") == "unknown")
        districts.append(
            {
                "district": district,
                "tasks": tasks,
                "total": len(tasks),
                "done": done,
                "warning": warning,
                "behind": behind,
                "unknown": unknown,
                "completion_avg": round(sum(rates) / len(rates), 2) if rates else None,
                "actual_total": sum(clean_number(item.get("actual_value")) or 0 for item in tasks),
            }
        )
    districts.sort(key=lambda item: (item["completion_avg"] is None, -(item["completion_avg"] or 0), item["district"]))

    total = sum(item["total"] for item in districts)
    done_count = sum(item["done"] for item in districts)
    warning_count = sum(item["warning"] for item in districts)
    behind_count = sum(item["behind"] for item in districts)
    unknown_count = sum(item["unknown"] for item in districts)
    done_rate = round(done_count / total * 100, 2) if total else 0
    return {
        "month": period["month"],
        "end_month": period["month"],
        "start_month": period["start_month"],
        "baseline_month": period["baseline_month"],
        "time_scope": period["time_scope"],
        "label": period["label"],
        "city": city,
        "districts": districts,
        "total": total,
        "done": done_count,
        "warning": warning_count,
        "behind": behind_count,
        "unknown": unknown_count,
        "done_rate": done_rate,
    }


def aggregate_for_period(month: str, start_month: str | None = None, time_scope: str | None = None) -> dict[str, Any]:
    rows, period = metric_rows_for_period(month, start_month, time_scope)
    return aggregate_from_rows(rows, period)


def aggregate_for_month(month: str) -> dict[str, Any]:
    return aggregate_for_period(month)


def build_report(month: str, period: str = "month", scope: str = "cumulative", start_month: str | None = None, time_scope: str | None = None) -> str:
    overview = aggregate_for_period(month, start_month, time_scope)
    label = overview["label"]
    lines = [
        f"# {label}指标完成情况报告",
        "",
        f"市级指标{len(overview['city'])}项；区级任务{overview['total']}项，已完成{overview['done']}项，接近完成{overview['warning']}项，未完成{overview['behind']}项。",
        "",
        "## 一、全市五项指标",
        "",
        "| 指标 | 目标 | 完成 | 完成率 | 状态 |",
        "| --- | ---: | ---: | ---: | --- |",
    ]
    for item in overview["city"]:
        lines.append(
            f"| {item['indicator_name']} | {format_value(item['target_value'], item['unit'])} | "
            f"{format_value(item['actual_value'], item['unit'])} | {format_value(item['completion_rate'], '%')} | {STATUS_LABELS.get(item['status'], item['status'])} |"
        )
    lines.extend(["", "## 二、各区两项任务", "", "| 区域 | 新增取证完成率 | 高级工以上完成率 |", "| --- | ---: | ---: |"])
    for district in overview["districts"]:
        task_map = {task["indicator_name"]: task for task in district["tasks"]}
        total_task = task_map.get(FORMAL_INDICATORS[0]["name"], {})
        advanced_task = task_map.get(FORMAL_INDICATORS[1]["name"], {})
        lines.append(f"| {district['district']} | {format_value(total_task.get('completion_rate'), '%')} | {format_value(advanced_task.get('completion_rate'), '%')} |")
    return "\n".join(lines)


def pdf_text(value: Any, style: Any) -> Any:
    from reportlab.platypus import Paragraph

    text = "-" if value is None or value == "" else str(value)
    return Paragraph(html.escape(text), style)


def pdf_display_district(name: Any) -> str:
    text = str(name or "").strip()
    if text in {"合计", "总计", "全市合计"}:
        return "全市"
    return text or "-"


def share_pdf_filename(period: dict[str, Any]) -> str:
    if period["time_scope"] == "range":
        stem = f"share_indicator_{period['start_month']}_to_{period['month']}"
    elif period["time_scope"] == "since2021":
        stem = f"share_indicator_since2021_to_{period['month']}"
    else:
        stem = f"share_indicator_{period['month']}"
    stem = re.sub(r"[^0-9A-Za-z._-]+", "_", stem).strip("_") or "share_indicator"
    return f"{stem}.pdf"


def make_pdf_table(data: list[list[Any]], widths: list[float]) -> Any:
    from reportlab.lib import colors
    from reportlab.platypus import Table, TableStyle

    table = Table(data, colWidths=widths, repeatRows=1, hAlign="LEFT")
    commands = [
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#123f39")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
        ("FONTSIZE", (0, 0), (-1, 0), 8.5),
        ("FONTSIZE", (0, 1), (-1, -1), 8),
        ("LEADING", (0, 0), (-1, -1), 10),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#d4dfda")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]
    for row_index in range(1, len(data)):
        if row_index % 2 == 0:
            commands.append(("BACKGROUND", (0, row_index), (-1, row_index), colors.HexColor("#f4f8f6")))
    table.setStyle(TableStyle(commands))
    return table


def build_share_pdf(month: str, start_month: str | None = None, time_scope: str | None = None) -> tuple[bytes, str]:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.platypus import SimpleDocTemplate, Spacer, Table, TableStyle
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"当前环境缺少 PDF 生成组件：{exc}") from exc

    if "STSong-Light" not in pdfmetrics.getRegisteredFontNames():
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))

    overview = aggregate_for_period(month, start_month, time_scope)
    ai_rows, _ = ai_district_rows_for_period(month, start_month, time_scope)
    level_rows, _ = certificate_level_rows_for_period(month, start_month, time_scope)
    total_level = next((row for row in level_rows if row.get("district") == "合计"), None) or {}
    city_rows = overview.get("city") or []
    city_done = sum(1 for item in city_rows if item.get("status") == "done")
    city_warning = sum(1 for item in city_rows if item.get("status") == "warning")
    city_behind = sum(1 for item in city_rows if item.get("status") == "behind")
    city_unknown = sum(1 for item in city_rows if item.get("status") == "unknown")

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        leftMargin=22,
        rightMargin=22,
        topMargin=22,
        bottomMargin=22,
        title=f"{overview['label']}指标完成情况分享页",
    )
    sample = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ShareTitle",
        parent=sample["Title"],
        fontName="STSong-Light",
        fontSize=20,
        leading=24,
        textColor=colors.HexColor("#123f39"),
        wordWrap="CJK",
        spaceAfter=8,
    )
    section_style = ParagraphStyle(
        "ShareSection",
        parent=sample["Heading2"],
        fontName="STSong-Light",
        fontSize=12,
        leading=15,
        textColor=colors.HexColor("#123f39"),
        wordWrap="CJK",
        spaceBefore=8,
        spaceAfter=5,
    )
    body_style = ParagraphStyle(
        "ShareBody",
        parent=sample["BodyText"],
        fontName="STSong-Light",
        fontSize=8.2,
        leading=10.5,
        wordWrap="CJK",
    )
    small_style = ParagraphStyle(
        "ShareSmall",
        parent=body_style,
        fontSize=7.4,
        leading=9,
    )

    story: list[Any] = [
        pdf_text(f"{overview['label']}指标完成情况分享页", title_style),
        pdf_text(
            f"市级{len(city_rows)}项指标：已完成{city_done}项、接近完成{city_warning}项、未完成{city_behind}项、待判断{city_unknown}项；"
            f"区级{overview.get('total') or 0}项任务，覆盖{len(overview.get('districts') or [])}个区域口径。",
            body_style,
        ),
        Spacer(1, 8),
    ]

    summary_data = [
        [pdf_text("市级指标", body_style), pdf_text("已完成", body_style), pdf_text("接近完成", body_style), pdf_text("未完成", body_style)],
        [pdf_text(len(city_rows), body_style), pdf_text(city_done, body_style), pdf_text(city_warning, body_style), pdf_text(city_behind, body_style)],
    ]
    summary_table = Table(summary_data, colWidths=[doc.width / 4] * 4, hAlign="LEFT")
    summary_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8f3ef")),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#123f39")),
                ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#d4dfda")),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.extend([summary_table, Spacer(1, 7), pdf_text("一、全市五项指标", section_style)])

    city_data = [[pdf_text("指标", body_style), pdf_text("目标", body_style), pdf_text("完成", body_style), pdf_text("完成率", body_style), pdf_text("状态", body_style)]]
    for item in city_rows:
        city_data.append(
            [
                pdf_text(item.get("indicator_name"), small_style),
                pdf_text(format_value(item.get("target_value"), item.get("unit")), small_style),
                pdf_text(format_value(item.get("actual_value"), item.get("unit")), small_style),
                pdf_text(format_value(item.get("completion_rate"), "%"), small_style),
                pdf_text(STATUS_LABELS.get(item.get("status"), item.get("status")), small_style),
            ]
        )
    story.append(make_pdf_table(city_data, [doc.width * 0.43, doc.width * 0.13, doc.width * 0.14, doc.width * 0.13, doc.width * 0.17]))

    story.extend([Spacer(1, 7), pdf_text("二、各区两项任务", section_style)])
    district_data = [
        [
            pdf_text("区域", body_style),
            pdf_text("新增取证目标", body_style),
            pdf_text("新增取证完成", body_style),
            pdf_text("完成率", body_style),
            pdf_text("高级以上目标", body_style),
            pdf_text("高级以上完成", body_style),
            pdf_text("完成率", body_style),
        ]
    ]
    for district in overview.get("districts") or []:
        task_map = {task.get("indicator_name"): task for task in district.get("tasks") or []}
        total_task = task_map.get(FORMAL_INDICATORS[0]["name"], {})
        advanced_task = task_map.get(FORMAL_INDICATORS[1]["name"], {})
        district_data.append(
            [
                pdf_text(pdf_display_district(district.get("district")), small_style),
                pdf_text(format_value(total_task.get("target_value"), total_task.get("unit")), small_style),
                pdf_text(format_value(total_task.get("actual_value"), total_task.get("unit")), small_style),
                pdf_text(format_value(total_task.get("completion_rate"), "%"), small_style),
                pdf_text(format_value(advanced_task.get("target_value"), advanced_task.get("unit")), small_style),
                pdf_text(format_value(advanced_task.get("actual_value"), advanced_task.get("unit")), small_style),
                pdf_text(format_value(advanced_task.get("completion_rate"), "%"), small_style),
            ]
        )
    story.append(make_pdf_table(district_data, [doc.width * 0.11, doc.width * 0.15, doc.width * 0.15, doc.width * 0.11, doc.width * 0.16, doc.width * 0.16, doc.width * 0.16]))

    story.extend([Spacer(1, 7), pdf_text("三、技师专项统计", section_style)])
    level_data = [
        [pdf_text("口径", body_style), pdf_text("特级/特技技师", body_style), pdf_text("首席技师", body_style)],
        [
            pdf_text("全市累计", small_style),
            pdf_text(format_value(total_level.get("special_technician"), total_level.get("unit") or "人次"), small_style),
            pdf_text(format_value(total_level.get("chief_technician"), total_level.get("unit") or "人次"), small_style),
        ],
    ]
    for item in [row for row in level_rows if row.get("district") != "合计"][:12]:
        level_data.append(
            [
                pdf_text(pdf_display_district(item.get("district")), small_style),
                pdf_text(format_value(item.get("special_technician"), item.get("unit") or "人次"), small_style),
                pdf_text(format_value(item.get("chief_technician"), item.get("unit") or "人次"), small_style),
            ]
        )
    story.append(make_pdf_table(level_data, [doc.width * 0.34, doc.width * 0.33, doc.width * 0.33]))

    story.extend([Spacer(1, 7), pdf_text("四、各区人工智能类证书", section_style)])
    ai_data = [[pdf_text("区域", body_style), pdf_text("完成人次", body_style), pdf_text("来源", body_style)]]
    for item in ai_rows:
        ai_data.append(
            [
                pdf_text(pdf_display_district(item.get("district")), small_style),
                pdf_text(format_value(item.get("actual_value"), item.get("unit") or "人次"), small_style),
                pdf_text(item.get("source_sheet") or "-", small_style),
            ]
        )
    story.append(make_pdf_table(ai_data, [doc.width * 0.22, doc.width * 0.18, doc.width * 0.6]))

    story.append(Spacer(1, 6))
    story.append(pdf_text("本 PDF 由只读分享页生成，仅展示指标完成情况，不包含上传、修改、删除权限。", small_style))

    doc.build(story)
    return buffer.getvalue(), share_pdf_filename(overview)


def build_word_report(
    month: str,
    period: str = "month",
    scope: str = "cumulative",
    kind: str = "indicator_report",
    start_month: str | None = None,
    time_scope: str | None = None,
) -> Path:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    period_info = resolve_time_period(month, start_month, time_scope)
    range_part = f"{period_info['start_month']}_to_{period_info['month']}" if period_info["time_scope"] == "range" else period_info["month"]
    suffix = f"{range_part}_{normalize_report_period(period)}_{normalize_report_scope(scope)}.docx"
    path = REPORT_DIR / f"{kind}_{suffix}"
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"当前环境缺少 python-docx，无法生成 Word：{exc}") from exc

    overview = aggregate_for_period(month, start_month, time_scope)
    doc = Document()
    doc.styles["Normal"].font.name = "Microsoft YaHei"
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.add_heading(f"{overview['label']}指标完成情况", level=1)
    doc.add_paragraph(f"市级指标{len(overview['city'])}项；区级任务{overview['total']}项，已完成{overview['done']}项，接近完成{overview['warning']}项，未完成{overview['behind']}项。")

    doc.add_heading("一、全市五项指标", level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for idx, header in enumerate(["指标", "目标", "完成", "完成率", "状态"]):
        table.rows[0].cells[idx].text = header
    for item in overview["city"]:
        cells = table.add_row().cells
        cells[0].text = item["indicator_name"]
        cells[1].text = format_value(item["target_value"], item["unit"])
        cells[2].text = format_value(item["actual_value"], item["unit"])
        cells[3].text = format_value(item["completion_rate"], "%")
        cells[4].text = STATUS_LABELS.get(item["status"], item["status"])

    doc.add_heading("二、各区两项任务", level=2)
    district_table = doc.add_table(rows=1, cols=7)
    district_table.style = "Table Grid"
    headers = ["区域", "新增取证目标", "新增取证完成", "完成率", "高级工以上目标", "高级工以上完成", "完成率"]
    for idx, header in enumerate(headers):
        district_table.rows[0].cells[idx].text = header
    for district in overview["districts"]:
        task_map = {task["indicator_name"]: task for task in district["tasks"]}
        total_task = task_map.get(FORMAL_INDICATORS[0]["name"], {})
        advanced_task = task_map.get(FORMAL_INDICATORS[1]["name"], {})
        cells = district_table.add_row().cells
        values = [
            district["district"],
            format_value(total_task.get("target_value"), total_task.get("unit")),
            format_value(total_task.get("actual_value"), total_task.get("unit")),
            format_value(total_task.get("completion_rate"), "%"),
            format_value(advanced_task.get("target_value"), advanced_task.get("unit")),
            format_value(advanced_task.get("actual_value"), advanced_task.get("unit")),
            format_value(advanced_task.get("completion_rate"), "%"),
        ]
        for idx, value in enumerate(values):
            cells[idx].text = value
    doc.add_paragraph("说明：本报告由平台根据上传表格和既定指标口径自动生成。")
    doc.save(path)
    return path


def is_local_admin_request(request: Request) -> bool:
    client_host = request.client.host if request.client else ""
    host_header = request.headers.get("host", "").split(":")[0].strip("[]").lower()
    local_hosts = {"localhost", "127.0.0.1", "::1"}
    local_client = client_host in {"127.0.0.1", "::1"} or client_host.startswith("127.")
    return local_client and host_header in local_hosts


def configured_admin_password() -> str:
    return os.getenv("INDICATOR_ADMIN_PASSWORD", "").strip()


def configured_admin_user() -> str:
    return os.getenv("INDICATOR_ADMIN_USER", "admin").strip() or "admin"


def is_password_admin_request(request: Request) -> bool:
    password = configured_admin_password()
    if not password:
        return False
    auth = request.headers.get("authorization", "")
    scheme, _, value = auth.partition(" ")
    if scheme.lower() != "basic" or not value:
        return False
    try:
        decoded = base64.b64decode(value).decode("utf-8")
    except (ValueError, UnicodeDecodeError):
        return False
    username, _, provided_password = decoded.partition(":")
    return secrets.compare_digest(username, configured_admin_user()) and secrets.compare_digest(provided_password, password)


def is_admin_request(request: Request) -> bool:
    return is_local_admin_request(request) or is_password_admin_request(request)


def admin_auth_challenge(detail: str = "请输入后台管理账号密码") -> None:
    raise HTTPException(
        status_code=401,
        detail=detail,
        headers={"WWW-Authenticate": 'Basic realm="Indicator Platform Admin", charset="UTF-8"'},
    )


def require_admin(request: Request) -> None:
    if is_admin_request(request):
        return
    if configured_admin_password():
        admin_auth_challenge()
    raise HTTPException(status_code=403, detail="后台管理密码未配置，公网仅开放只读分享页")


def require_local_admin(request: Request) -> None:
    if not is_local_admin_request(request):
        raise HTTPException(status_code=403, detail="该操作只能在服务器本机执行")


def local_ipv4_addresses() -> list[str]:
    addresses: list[str] = []
    try:
        with socket.socket(socket.AF_INET, socket.SOCK_DGRAM) as sock:
            sock.connect(("223.5.5.5", 80))
            addresses.append(sock.getsockname()[0])
    except OSError:
        pass
    try:
        for item in socket.gethostbyname_ex(socket.gethostname())[2]:
            if item and not item.startswith("127."):
                addresses.append(item)
    except OSError:
        pass
    unique: list[str] = []
    for address in addresses:
        if address not in unique:
            unique.append(address)
    return unique


app = FastAPI(title=APP_TITLE)
app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")


@app.middleware("http")
async def no_cache_for_pages_and_assets(request: Request, call_next: Any):
    response = await call_next(request)
    path = request.url.path
    if path in {"/", "/share"} or path.startswith("/static/"):
        response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
    return response


@app.on_event("startup")
def on_startup() -> None:
    init_db()
    backfill_certificate_level_metrics()


@app.get("/", response_class=HTMLResponse)
def index(request: Request):
    if not is_local_admin_request(request):
        return RedirectResponse(url="/share", status_code=307)
    return FileResponse(STATIC_DIR / "index.html")


@app.get("/admin", response_class=HTMLResponse)
def admin_page(request: Request) -> FileResponse:
    if not is_admin_request(request):
        if configured_admin_password():
            admin_auth_challenge()
        raise HTTPException(status_code=403, detail="后台管理密码未配置")
    return FileResponse(STATIC_DIR / "index.html", headers={"Cache-Control": "no-store"})


@app.get("/share", response_class=HTMLResponse)
def share_page() -> FileResponse:
    return FileResponse(STATIC_DIR / "share.html", headers={"Cache-Control": "no-store"})


@app.get("/api/health")
def health() -> dict[str, Any]:
    return {"ok": True, "title": APP_TITLE, "month": latest_month()}


@app.get("/api/share-info")
def share_info(request: Request) -> dict[str, Any]:
    port = request.url.port
    default_port = 443 if request.url.scheme == "https" else 80
    port_part = "" if not port or port == default_port else f":{port}"
    return {
        "admin_local": is_local_admin_request(request),
        "preview_url": str(request.url_for("share_page")),
        "urls": [f"{request.url.scheme}://{address}{port_part}/share" for address in local_ipv4_addresses()],
    }


def read_public_share_url() -> str:
    if GITHUB_PAGES_SHARE_URL:
        return GITHUB_PAGES_SHARE_URL
    if not PUBLIC_SHARE_URL_FILE.exists():
        return ""
    text = PUBLIC_SHARE_URL_FILE.read_text(encoding="utf-8", errors="ignore")
    match = re.search(r"https://[^\s]+", text)
    return match.group(0).rstrip("。.,，") if match else ""


def is_github_pages_share_url(url: str) -> bool:
    return bool(GITHUB_PAGES_SHARE_URL and url.rstrip("/") == GITHUB_PAGES_SHARE_URL.rstrip("/"))


def read_public_share_status(url: str) -> dict[str, Any]:
    if not url or not PUBLIC_SHARE_STATUS_FILE.exists():
        return {}
    try:
        payload = json.loads(PUBLIC_SHARE_STATUS_FILE.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError):
        return {}
    if payload.get("url") != url:
        return {}
    return payload


def public_share_note(status: dict[str, Any], url: str) -> str:
    if is_github_pages_share_url(url):
        checked = status.get("checked_at", "")
        if status.get("reachable") is False:
            detail = status.get("detail") or "请稍后重试，GitHub Pages 可能仍在构建或被网络拦截。"
            return f"GitHub Pages 固定公网链接暂不可访问：{detail}"
        return f"GitHub Pages 固定公网链接，可直接微信查看。最近检测时间：{checked}。" if checked else "GitHub Pages 固定公网链接，可直接微信查看。"
    if not url:
        return "尚未生成公网分享链接；如需长期固定访问，建议使用云服务器或 Cloudflare named tunnel。"
    if status.get("reachable") is True:
        checked = status.get("checked_at", "")
        if status.get("refreshed"):
            return f"旧链接失效后已自动更新，新公网链接可访问，最近检测时间：{checked}。"
        return f"公网链接可访问，最近检测时间：{checked}。" if checked else "公网链接可访问。"
    if status.get("reachable") is False:
        detail = status.get("detail") or "请重新生成公网链接，或改用固定域名/云服务器。"
        if status.get("refresh_attempted"):
            return f"公网链接已失效，自动更新失败：{detail}"
        return f"公网链接已失效或无法访问：{detail}"
    return "当前为临时公网链接，点击“检测/更新”确认是否可访问；若失效会重新生成链接。长期固定链接需要域名并配置 Cloudflare named tunnel 或云服务器。"


def public_share_payload(url: str, status: dict[str, Any] | None = None) -> dict[str, Any]:
    status = status or read_public_share_status(url)
    if is_github_pages_share_url(url):
        reachable = status.get("reachable") if "reachable" in status else True
        return {
            "url": url,
            "last_url": url,
            "fixed": True,
            "provider": "GitHub Pages",
            "reachable": reachable,
            "status": status.get("status", "ok" if reachable is not False else "invalid"),
            "checked_at": status.get("checked_at", ""),
            "detail": status.get("detail", ""),
            "refreshed": False,
            "refresh_attempted": False,
            "note": public_share_note(status, url),
        }
    if public_share_disabled():
        return {
            "url": "",
            "last_url": url,
            "fixed": False,
            "provider": "",
            "reachable": False,
            "status": "disabled",
            "checked_at": "",
            "detail": "公网自动分享已停用",
            "refreshed": False,
            "refresh_attempted": False,
            "note": "公网自动分享已停用；不会自动生成或复制外网链接。本机访问和内网页面不受影响。",
        }
    reachable = status.get("reachable") if "reachable" in status else None
    publishable_url = url if reachable is not False else ""
    return {
        "url": publishable_url,
        "last_url": url,
        "fixed": False,
        "provider": "Cloudflare quick tunnel" if url else "",
        "reachable": reachable,
        "status": status.get("status", "unchecked" if url else "missing"),
        "checked_at": status.get("checked_at", ""),
        "detail": status.get("detail", ""),
        "refreshed": bool(status.get("refreshed")),
        "refresh_attempted": bool(status.get("refresh_attempted")),
        "note": public_share_note(status, url),
    }


def save_public_share_status(status: dict[str, Any]) -> None:
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    PUBLIC_SHARE_STATUS_FILE.write_text(json.dumps(status, ensure_ascii=False, indent=2), encoding="utf-8")


def public_share_disabled() -> bool:
    return (DATA_DIR / "public-share-disabled.flag").exists()


def write_public_share_url(url: str) -> None:
    lines = [
        "Public read-only share URL:",
        url,
        "",
        "Notes:",
        "1. This link uses Cloudflare quick tunnel and does not show the localtunnel IP warning page.",
        "2. The public URL exposes only the read-only /share page. Admin write/download APIs remain blocked for public visitors.",
        "3. Quick tunnel URLs can change after the tunnel restarts. A permanently fixed URL requires a domain name plus a Cloudflare named tunnel, or a cloud server and domain.",
        "4. If the computer is powered off, sleeping, or offline, the public URL is unavailable until the machine is back online.",
        f"5. Last refreshed: {now_iso()}",
    ]
    PUBLIC_SHARE_URL_FILE.write_text("\n".join(lines) + "\n", encoding="utf-8")


def stop_cloudflared_public_tunnels() -> None:
    if os.name == "nt":
        command = (
            "Get-CimInstance Win32_Process | "
            "Where-Object { $_.Name -match 'cloudflared' -and $_.CommandLine -like '*127.0.0.1:8020*' } | "
            "ForEach-Object { try { Stop-Process -Id $_.ProcessId -Force } catch {} }"
        )
        subprocess.run(
            ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", command],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
            timeout=10,
            creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
        )
        return
    subprocess.run(
        ["pkill", "-f", "cloudflared.*127.0.0.1:8020"],
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        timeout=10,
    )


def extract_cloudflare_url_from_logs() -> str:
    pattern = re.compile(r"https://[a-zA-Z0-9-]+\.trycloudflare\.com")
    for path in (CLOUDFLARED_LOG, CLOUDFLARED_ERR_LOG, CLOUDFLARED_OUT_LOG):
        if not path.exists():
            continue
        text = path.read_text(encoding="utf-8", errors="ignore")
        match = pattern.search(text)
        if match:
            return f"{match.group(0)}/share"
    return ""


def cloudflared_tunnel_ready() -> bool:
    for path in (CLOUDFLARED_LOG, CLOUDFLARED_ERR_LOG):
        if path.exists() and "Registered tunnel connection" in path.read_text(encoding="utf-8", errors="ignore"):
            return True
    return False


def local_share_ready() -> bool:
    status = probe_public_share_url("http://127.0.0.1:8020/share")
    return status.get("reachable") is True


def start_cloudflared_tunnel_once() -> tuple[str, str]:
    for path in (CLOUDFLARED_LOG, CLOUDFLARED_OUT_LOG, CLOUDFLARED_ERR_LOG):
        try:
            path.unlink(missing_ok=True)
        except OSError:
            pass
    args = [
        str(CLOUDFLARED_EXE),
        "tunnel",
        "--url",
        "http://127.0.0.1:8020",
        "--protocol",
        "http2",
        "--logfile",
        str(CLOUDFLARED_LOG),
        "--loglevel",
        "info",
    ]
    try:
        with CLOUDFLARED_OUT_LOG.open("ab") as stdout, CLOUDFLARED_ERR_LOG.open("ab") as stderr:
            subprocess.Popen(
                args,
                cwd=str(ROOT),
                stdout=stdout,
                stderr=stderr,
                creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0),
            )
    except OSError as exc:
        return "", f"启动 Cloudflare quick tunnel 失败：{exc}"
    candidate_url = ""
    for _ in range(60):
        candidate_url = candidate_url or extract_cloudflare_url_from_logs()
        if candidate_url and cloudflared_tunnel_ready():
            return candidate_url, "Cloudflare quick tunnel 已连接"
        time.sleep(1)
    if candidate_url:
        return candidate_url, "Cloudflare quick tunnel 已生成地址，但未确认连接成功"
    return "", "Cloudflare quick tunnel 启动超时，未能获取新公网链接"


def refresh_public_share_url(max_attempts: int = 3) -> tuple[str, dict[str, Any]]:
    last_status: dict[str, Any] = {
        "url": "",
        "checked_at": now_iso(),
        "reachable": False,
        "status": "invalid",
        "detail": "",
        "refresh_attempted": True,
    }
    if not CLOUDFLARED_EXE.exists():
        last_status["detail"] = f"未找到 cloudflared：{CLOUDFLARED_EXE}"
        return "", last_status
    if not local_share_ready():
        last_status["detail"] = "本机 /share 页面不可访问，未启动公网隧道"
        return "", last_status
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    for attempt in range(1, max_attempts + 1):
        try:
            stop_cloudflared_public_tunnels()
        except (subprocess.SubprocessError, OSError) as exc:
            last_status["detail"] = f"停止旧公网隧道失败：{exc}"
            return "", last_status
        time.sleep(1)
        candidate_url, start_detail = start_cloudflared_tunnel_once()
        if not candidate_url:
            last_status["detail"] = f"第 {attempt} 次更新失败：{start_detail}"
            continue
        for probe_attempt in range(1, 7):
            status = probe_public_share_url(candidate_url)
            status["refreshed"] = True
            status["refresh_attempted"] = True
            status["refresh_detail"] = f"第 {attempt} 次更新：{start_detail}"
            if status.get("reachable") is True:
                write_public_share_url(candidate_url)
                return candidate_url, status
            last_status = status
            if probe_attempt < 6:
                time.sleep(2)
    if last_status.get("url"):
        last_status["detail"] = f"已尝试 {max_attempts} 次更新，但新链接仍不可访问：{last_status.get('detail', '')}"
    elif not last_status.get("detail"):
        last_status["detail"] = f"已尝试 {max_attempts} 次更新，但未能获取新公网链接"
    return "", last_status


def probe_public_share_url(url: str) -> dict[str, Any]:
    checked_at = now_iso()
    base = {
        "url": url,
        "checked_at": checked_at,
        "reachable": False,
        "status": "invalid",
        "detail": "",
    }
    if not url:
        base.update({"status": "missing", "detail": "尚未生成公网分享链接"})
        return base
    try:
        request = urllib.request.Request(
            url,
            headers={
                "User-Agent": "IndicatorPlatformLinkCheck/1.0",
                "Cache-Control": "no-cache",
            },
        )
        with urllib.request.urlopen(request, timeout=8) as response:
            code = response.getcode()
            content_type = response.headers.get("content-type", "")
            body = response.read(65536).decode("utf-8", errors="ignore")
    except urllib.error.HTTPError as exc:
        base["detail"] = f"公网返回 HTTP {exc.code}"
        return base
    except (urllib.error.URLError, TimeoutError, OSError) as exc:
        reason = getattr(exc, "reason", exc)
        base["detail"] = f"连接失败或超时：{reason}"
        return base
    is_html = "text/html" in content_type.lower() or "<html" in body[:500].lower()
    looks_like_share_page = (
        "指标完成情况平台" in body
        or "/static/share.js" in body
        or "share-root" in body
        or "指标完成情况公开只读页" in body
        or "GitHub Pages 静态只读版" in body
        or "data/site-data.json" in body
    )
    if 200 <= code < 400 and is_html and looks_like_share_page:
        base.update({"reachable": True, "status": "ok", "detail": f"HTTP {code}"})
        return base
    base["detail"] = f"公网返回 HTTP {code}，但不是平台分享页"
    return base


@app.get("/api/public-share-url")
def public_share_url() -> dict[str, Any]:
    url = read_public_share_url()
    return public_share_payload(url)


@app.post("/api/public-share-url/check")
def check_public_share_url(request: Request) -> dict[str, Any]:
    require_admin(request)
    url = read_public_share_url()
    if is_github_pages_share_url(url):
        status = probe_public_share_url(url)
        save_public_share_status(status)
        return public_share_payload(url, status)
    if public_share_disabled():
        status = {
            "url": url,
            "checked_at": now_iso(),
            "reachable": False,
            "status": "disabled",
            "detail": "公网自动分享已停用",
            "refresh_attempted": False,
        }
        save_public_share_status(status)
        return public_share_payload(url, status)
    status = probe_public_share_url(url)
    if status.get("reachable") is not True:
        new_url, refreshed_status = refresh_public_share_url()
        if new_url:
            url = new_url
            status = refreshed_status
        else:
            status = refreshed_status
            if url:
                status["url"] = url
    save_public_share_status(status)
    return public_share_payload(url, status)


@app.get("/api/meta")
def meta() -> dict[str, Any]:
    with db() as conn:
        months = [row["month"] for row in conn.execute("SELECT DISTINCT month FROM metrics").fetchall()]
        districts = [
            row["district"]
            for row in conn.execute(
                "SELECT DISTINCT district FROM metrics WHERE district != '合计' ORDER BY district"
            ).fetchall()
        ]
        categories = [row["category"] for row in conn.execute("SELECT DISTINCT category FROM metrics ORDER BY category").fetchall()]
    months = sorted(months, key=period_sort_value, reverse=True)
    return {
        "months": months,
        "districts": districts,
        "categories": categories,
        "statuses": [{"value": key, "label": label} for key, label in STATUS_LABELS.items()],
        "latest_month": months[0] if months else DEFAULT_MONTH,
    }


@app.get("/api/overview")
def overview(
    month: str | None = Query(None),
    start_month: str | None = Query(None),
    time_scope: str | None = Query(None),
) -> dict[str, Any]:
    return aggregate_for_period(normalize_month(month) if month else latest_month(), start_month, time_scope)


@app.get("/api/metrics")
def metrics(
    month: str | None = Query(None),
    start_month: str | None = Query(None),
    time_scope: str | None = Query(None),
    district: str | None = Query(None),
    category: str | None = Query(None),
    status: str | None = Query(None),
    q: str | None = Query(None),
    limit: int = Query(500, ge=1, le=5000),
) -> dict[str, Any]:
    target_month = normalize_month(month) if month else latest_month()
    rows, period = metric_rows_for_period(target_month, start_month, time_scope)
    filtered = filter_metric_rows(rows, district or None, category or None, status or None, q or None, limit)
    return {"items": filtered, "count": len(filtered), **period}


@app.get("/api/ai-districts")
def ai_districts(
    month: str | None = Query(None),
    start_month: str | None = Query(None),
    time_scope: str | None = Query(None),
) -> dict[str, Any]:
    target_month = normalize_month(month) if month else latest_month()
    rows, period = ai_district_rows_for_period(target_month, start_month, time_scope)
    return {"items": rows, "count": len(rows), **period}


@app.get("/api/level-stats")
def level_stats(
    month: str | None = Query(None),
    start_month: str | None = Query(None),
    time_scope: str | None = Query(None),
) -> dict[str, Any]:
    target_month = normalize_month(month) if month else latest_month()
    rows, period = certificate_level_rows_for_period(target_month, start_month, time_scope)
    total = next((row for row in rows if row["district"] == "合计"), None)
    return {"items": rows, "total": total, "count": len(rows), **period}


@app.get("/api/share/pdf")
def download_share_pdf(
    month: str | None = Query(None),
    start_month: str | None = Query(None),
    time_scope: str | None = Query(None),
    disposition: str = Query("attachment"),
) -> Response:
    target_month = normalize_month(month) if month else latest_month()
    content, filename = build_share_pdf(target_month, start_month, time_scope)
    content_disposition = "inline" if disposition == "inline" else "attachment"
    return Response(
        content=content,
        media_type="application/pdf",
        headers={"Content-Disposition": f'{content_disposition}; filename="{filename}"'},
    )


@app.post("/api/upload")
async def upload(request: Request, month: str = Form(DEFAULT_MONTH), file: UploadFile = File(...)) -> dict[str, Any]:
    require_admin(request)
    normalized_month = infer_month_from_text(file.filename) or normalize_month(month)
    safe_name = re.sub(r"[^0-9A-Za-z._\-\u4e00-\u9fff]+", "_", file.filename or "upload.xlsx")
    stored_path = UPLOAD_DIR / f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{safe_name}"
    with stored_path.open("wb") as handle:
        shutil.copyfileobj(file.file, handle)
    upload_id = create_upload_record(file.filename or safe_name, str(stored_path), normalized_month)
    try:
        rows, sheet_names, ai_rows, level_rows = parse_spreadsheet(stored_path, normalized_month, upload_id, file.filename or safe_name)
        if not rows and not ai_rows and not level_rows:
            update_upload_record(upload_id, 0, sheet_names, "empty")
            raise HTTPException(status_code=422, detail="没有识别到指标数据，请检查表格列名和地区列。")
        insert_metrics(rows)
        insert_ai_district_metrics(ai_rows)
        insert_certificate_level_metrics(level_rows)
        update_upload_record(upload_id, len(rows) + len(ai_rows) + len(level_rows), sheet_names)
    except HTTPException:
        raise
    except Exception as exc:
        update_upload_record(upload_id, 0, [], "failed")
        raise HTTPException(status_code=500, detail=f"解析失败：{exc}") from exc
    return {"ok": True, "upload_id": upload_id, "month": normalized_month, "rows": len(rows) + len(ai_rows) + len(level_rows), "sheets": sheet_names}


@app.put("/api/manual/professional-skill")
def save_professional_skill_metric(request: Request, payload: ManualMetricPayload) -> dict[str, Any]:
    require_admin(request)
    target_month = normalize_month(payload.month) if payload.month else latest_month()
    indicator = FORMAL_INDICATORS[3]
    if payload.actual_value is None:
        with db() as conn:
            conn.execute(
                "DELETE FROM metrics WHERE month = ? AND district = ? AND indicator_name = ?",
                (target_month, "合计", indicator["name"]),
            )
        return {"ok": True, "month": target_month, "cleared": True, "overview": aggregate_for_month(target_month)}
    actual = round(float(payload.actual_value), 2)
    if actual < 0:
        raise HTTPException(status_code=422, detail="手工填入数字不能小于0")
    insert_metrics([city_metric_row(indicator, actual, target_month, None, "手工填入", indicator["note"])])
    return {"ok": True, "month": target_month, "overview": aggregate_for_month(target_month)}


@app.get("/api/report")
def report(
    month: str | None = Query(None),
    period: str = Query("month"),
    scope: str = Query("cumulative"),
    start_month: str | None = Query(None),
    time_scope: str | None = Query(None),
) -> dict[str, Any]:
    target_month = normalize_month(month) if month else latest_month()
    period_info = resolve_time_period(target_month, start_month, time_scope)
    return {"period": normalize_report_period(period), "scope": normalize_report_scope(scope), "markdown": build_report(target_month, period, scope, start_month, time_scope), **period_info}


@app.get("/api/report/download")
def download_report(
    request: Request,
    month: str | None = Query(None),
    period: str = Query("month"),
    scope: str = Query("cumulative"),
    start_month: str | None = Query(None),
    time_scope: str | None = Query(None),
) -> PlainTextResponse:
    require_admin(request)
    target_month = normalize_month(month) if month else latest_month()
    period_info = resolve_time_period(target_month, start_month, time_scope)
    markdown = build_report(target_month, period, scope, start_month, time_scope)
    range_part = f"{period_info['start_month']}_to_{period_info['month']}" if period_info["time_scope"] == "range" else period_info["month"]
    filename = f"indicator_report_{range_part}_{normalize_report_period(period)}_{normalize_report_scope(scope)}.md"
    return PlainTextResponse(markdown, media_type="text/markdown; charset=utf-8", headers={"Content-Disposition": f'attachment; filename="{filename}"'})


@app.get("/api/report/word")
def download_word_report(
    request: Request,
    month: str | None = Query(None),
    period: str = Query("month"),
    scope: str = Query("cumulative"),
    start_month: str | None = Query(None),
    time_scope: str | None = Query(None),
) -> FileResponse:
    require_admin(request)
    target_month = normalize_month(month) if month else latest_month()
    report_path = build_word_report(target_month, period, scope, "indicator_report", start_month, time_scope)
    return FileResponse(report_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=report_path.name)


@app.get("/api/report/monthly-word")
def download_monthly_progress_word(
    request: Request,
    month: str | None = Query(None),
    period: str = Query("month"),
    scope: str = Query("cumulative"),
    start_month: str | None = Query(None),
    time_scope: str | None = Query(None),
) -> FileResponse:
    require_admin(request)
    target_month = normalize_month(month) if month else latest_month()
    report_path = build_word_report(target_month, period, scope, "monthly_progress", start_month, time_scope)
    return FileResponse(report_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=report_path.name)


@app.get("/api/report/quarter-word")
def download_quarter_district_word(
    request: Request,
    month: str | None = Query(None),
    period: str = Query("quarter"),
    scope: str = Query("cumulative"),
    start_month: str | None = Query(None),
    time_scope: str | None = Query(None),
) -> FileResponse:
    require_admin(request)
    target_month = normalize_month(month) if month else latest_month()
    report_path = build_word_report(target_month, period, scope, "quarter_district", start_month, time_scope)
    return FileResponse(report_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=report_path.name)


def open_in_file_manager(target: Path) -> None:
    resolved = target.resolve()
    if sys.platform.startswith("win"):
        if resolved.is_file():
            subprocess.Popen(f'explorer.exe /select,"{resolved}"')
        else:
            subprocess.Popen(f'explorer.exe "{resolved}"')
        return
    folder = resolved.parent if resolved.is_file() else resolved
    command = ["open", "-R", str(resolved)] if sys.platform == "darwin" and resolved.is_file() else ["xdg-open", str(folder)]
    subprocess.Popen(command)


@app.post("/api/report/open-folder")
def open_report_folder(request: Request, payload: OpenReportFolderPayload) -> dict[str, Any]:
    require_local_admin(request)
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    filename = Path(payload.filename or "").name
    target = REPORT_DIR / filename if filename else REPORT_DIR
    if filename and not target.exists():
        target = REPORT_DIR
    open_in_file_manager(target)
    selected = target if target.exists() else REPORT_DIR
    folder = selected.parent if selected.is_file() else selected
    return {"ok": True, "folder": str(folder), "path": str(selected)}


@app.get("/api/uploads")
def uploads(request: Request) -> dict[str, Any]:
    require_admin(request)
    with db() as conn:
        rows = rows_to_dicts(conn.execute("SELECT * FROM uploads ORDER BY uploaded_at DESC, id DESC LIMIT 50").fetchall())
    for row in rows:
        try:
            row["sheet_names"] = json.loads(row.get("sheet_names") or "[]")
        except json.JSONDecodeError:
            row["sheet_names"] = []
    return {"items": rows}


@app.patch("/api/uploads/{upload_id}")
def update_upload(request: Request, upload_id: int, payload: UploadUpdatePayload) -> dict[str, Any]:
    require_admin(request)
    filename = (payload.filename or "").strip()
    if not filename:
        raise HTTPException(status_code=422, detail="文件名不能为空")
    if len(filename) > 180:
        raise HTTPException(status_code=422, detail="文件名不能超过180个字符")
    with db() as conn:
        row = conn.execute("SELECT * FROM uploads WHERE id = ?", (upload_id,)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="上传记录不存在")
        conn.execute("UPDATE uploads SET filename = ? WHERE id = ?", (filename, upload_id))
        updated = conn.execute("SELECT * FROM uploads WHERE id = ?", (upload_id,)).fetchone()
    item = dict(updated)
    try:
        item["sheet_names"] = json.loads(item.get("sheet_names") or "[]")
    except json.JSONDecodeError:
        item["sheet_names"] = []
    return {"ok": True, "item": item}


@app.delete("/api/uploads/{upload_id}")
def delete_upload(request: Request, upload_id: int) -> dict[str, Any]:
    require_admin(request)
    with db() as conn:
        row = conn.execute("SELECT * FROM uploads WHERE id = ?", (upload_id,)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="上传记录不存在")
        metric_count = conn.execute("SELECT COUNT(*) AS c FROM metrics WHERE upload_id = ?", (upload_id,)).fetchone()["c"]
        conn.execute("DELETE FROM ai_district_metrics WHERE upload_id = ?", (upload_id,))
        conn.execute("DELETE FROM certificate_level_metrics WHERE upload_id = ?", (upload_id,))
        conn.execute("DELETE FROM metrics WHERE upload_id = ?", (upload_id,))
        conn.execute("DELETE FROM uploads WHERE id = ?", (upload_id,))
    stored_path = str(row["stored_path"] or "").strip()
    if stored_path:
        try:
            path = Path(stored_path)
            if path.exists() and path.is_file() and path.parent.resolve() == UPLOAD_DIR.resolve():
                path.unlink()
        except OSError:
            pass
    return {"ok": True, "upload_id": upload_id, "deleted_metrics": metric_count, "filename": row["filename"]}


@app.get("/api/rules")
def list_rules(request: Request) -> dict[str, Any]:
    require_admin(request)
    with db() as conn:
        rows = rows_to_dicts(conn.execute("SELECT * FROM indicator_rules ORDER BY id DESC").fetchall())
    return {"items": rows}


@app.post("/api/rules")
def create_rule(request: Request, payload: RulePayload) -> dict[str, Any]:
    require_admin(request)
    if payload.direction not in {"higher_better", "lower_better"}:
        raise HTTPException(status_code=400, detail="direction 仅支持 higher_better 或 lower_better")
    with db() as conn:
        cursor = conn.execute(
            """
            INSERT INTO indicator_rules (indicator_pattern, category, direction, warning_threshold, done_threshold, created_at)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            (payload.indicator_pattern, payload.category, payload.direction, payload.warning_threshold, payload.done_threshold, now_iso()),
        )
        row = conn.execute("SELECT * FROM indicator_rules WHERE id = ?", (cursor.lastrowid,)).fetchone()
    return {"ok": True, "item": dict(row)}
